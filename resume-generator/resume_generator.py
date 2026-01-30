#!/usr/bin/env python3
"""Resume Generator - Creates professional resumes using Claude Agent SDK with custom MCP tools."""

import sys
from pathlib import Path
import json
from typing import Any

# Add shared module to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from claude_agent_sdk import (
    AssistantMessage,
    ClaudeSDKClient,
    ClaudeAgentOptions,
    HookContext,
    HookMatcher,
    ResultMessage,
    TextBlock,
    tool,
    create_sdk_mcp_server,
)
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from shared.cli import AgentCLI, CLIArgument, run_agent_cli
from mock_data import get_search_results, get_linkedin_profile as _get_linkedin_profile


# =============================================================================
# MCP Tool Definitions
# =============================================================================

@tool("web_search", "Search the web for information about a person", {"query": str})
async def web_search(args: dict) -> dict:
    """Search the web for information about a person."""
    query = args["query"]
    results = get_search_results(query)
    return {
        "content": [{"type": "text", "text": json.dumps(results, indent=2)}]
    }


@tool("get_linkedin_profile", "Fetch a person's LinkedIn profile", {"name": str})
async def get_linkedin_profile(args: dict) -> dict:
    """Fetch a person's LinkedIn profile data."""
    name = args["name"]
    profile = _get_linkedin_profile(name)
    return {
        "content": [{"type": "text", "text": json.dumps(profile, indent=2)}]
    }


@tool(
    "analyze_resume_data",
    "Analyze profile data and structure it for resume generation",
    {"name": str, "profile_json": str}
)
async def analyze_resume_data(args: dict) -> dict:
    """Analyze and structure profile data for resume generation."""
    name, profile_json = args["name"], args["profile_json"]
    profile = json.loads(profile_json)

    resume_data = {
        "header": {
            "name": profile.get("name", name),
            "headline": profile.get("headline", ""),
            "email": profile.get("email", ""),
            "phone": profile.get("phone", ""),
            "location": profile.get("location", ""),
            "linkedin": profile.get("linkedin_url", "")
        },
        "summary": profile.get("summary", ""),
        "experience": profile.get("experience", [])[:3],
        "education": profile.get("education", []),
        "skills": profile.get("skills", [])[:10]
    }

    return {
        "content": [{"type": "text", "text": json.dumps(resume_data, indent=2)}]
    }


@tool(
    "generate_resume_docx",
    "Generate a professional resume as a Word document",
    {"resume_json": str, "output_path": str}
)
async def generate_resume_docx(args: dict) -> dict:
    """Generate a resume document using python-docx."""
    resume_json, output_path = args["resume_json"], args["output_path"]
    resume = json.loads(resume_json)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set margins (0.5 inches)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header - Name
    header = resume.get("header", {})
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(header.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info
    contact = f"{header.get('email', '')} | {header.get('phone', '')} | {header.get('location', '')}"
    contact_para = doc.add_paragraph(contact)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(resume.get("summary", ""))

    # Experience
    doc.add_heading("Experience", level=1)
    for job in resume.get("experience", []):
        job_para = doc.add_paragraph()
        title_run = job_para.add_run(f"{job.get('title', '')} - {job.get('company', '')}")
        title_run.bold = True
        doc.add_paragraph(job.get("duration", ""), style="Normal")
        for bullet in job.get("description", []):
            doc.add_paragraph(bullet, style="List Bullet")

    # Education
    doc.add_heading("Education", level=1)
    for edu in resume.get("education", []):
        doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('school', '')} ({edu.get('year', '')})")

    # Skills
    doc.add_heading("Skills", level=1)
    skills = ", ".join(resume.get("skills", []))
    doc.add_paragraph(skills)

    doc.save(output_path)

    return {
        "content": [{"type": "text", "text": f"Resume successfully saved to {output_path}"}]
    }

@tool(
    "evaluate_resume_quality",
    "Evaluate a resume against quality criteria and return a score with feedback",
    {"resume_json": str}
)
async def evaluate_resume_quality(args: dict) -> dict:
    """
    Evaluate resume quality against 4 criteria (25 points each, total 100):
    1. Completeness: All sections present
    2. Conciseness: Summary < 3 sentences, bullets < 100 chars
    3. Specificity: Contains quantified achievements
    4. Formatting: Proper structure, no empty sections
    """
    resume = json.loads(args["resume_json"])
    score = 0
    issues = []
    suggestions = []

    # 1. Completeness check (25 points)
    required = ["header", "summary", "experience", "education", "skills"]
    missing = [s for s in required if s not in resume or not resume[s]]
    completeness_score = max(0, 25 - (5 * len(missing)))
    score += completeness_score
    if missing:
        issues.append(f"Missing sections: {', '.join(missing)}")
        suggestions.append(f"Add the following sections: {', '.join(missing)}")

    # 2. Conciseness check (25 points)
    conciseness_score = 25
    summary = resume.get("summary", "")
    sentence_count = len([s for s in summary.split('.') if s.strip()])
    if sentence_count > 3:
        conciseness_score -= 10
        issues.append(f"Summary has {sentence_count} sentences (should be ≤3)")
        suggestions.append("Shorten summary to 2-3 impactful sentences")

    # Check bullet point lengths
    long_bullets = []
    for job in resume.get("experience", []):
        for bullet in job.get("description", []):
            if len(bullet) > 100:
                long_bullets.append(bullet[:50] + "...")
    if long_bullets:
        conciseness_score -= min(15, len(long_bullets) * 5)
        issues.append(f"{len(long_bullets)} bullet points exceed 100 characters")
        suggestions.append("Shorten bullet points to under 100 characters each")
    score += max(0, conciseness_score)

    # 3. Specificity check (25 points) - look for metrics/numbers
    specificity_score = 0
    metrics_found = 0
    text_to_check = summary
    for job in resume.get("experience", []):
        text_to_check += " ".join(job.get("description", []))

    # Check for numbers, percentages, dollar amounts
    import re
    metrics_patterns = [
        r'\d+%',           # Percentages
        r'\$[\d,]+',       # Dollar amounts
        r'\d+\+?\s*(years|months|people|team|users|customers)',  # Quantities
        r'(increased|decreased|improved|reduced|grew|saved).*\d+',  # Impact metrics
    ]
    for pattern in metrics_patterns:
        matches = re.findall(pattern, text_to_check, re.IGNORECASE)
        metrics_found += len(matches)

    if metrics_found >= 5:
        specificity_score = 25
    elif metrics_found >= 3:
        specificity_score = 20
    elif metrics_found >= 1:
        specificity_score = 15
    else:
        specificity_score = 5
        issues.append("Resume lacks quantified achievements")
        suggestions.append("Add metrics: percentages, dollar amounts, team sizes, timeframes")
    score += specificity_score

    # 4. Formatting check (25 points)
    formatting_score = 25

    # Check header completeness
    header = resume.get("header", {})
    header_fields = ["name", "email", "location"]
    missing_header = [f for f in header_fields if not header.get(f)]
    if missing_header:
        formatting_score -= len(missing_header) * 5
        issues.append(f"Header missing: {', '.join(missing_header)}")

    # Check experience structure
    for job in resume.get("experience", []):
        if not job.get("title") or not job.get("company"):
            formatting_score -= 5
            issues.append("Experience entries missing title or company")
            break
        if not job.get("description"):
            formatting_score -= 5
            issues.append("Experience entries missing bullet points")
            break

    score += max(0, formatting_score)

    return {
        "content": [{
            "type": "text",
            "text": json.dumps({
                "score": score,
                "max_score": 100,
                "passing": score >= 75,
                "breakdown": {
                    "completeness": completeness_score,
                    "conciseness": max(0, conciseness_score),
                    "specificity": specificity_score,
                    "formatting": max(0, formatting_score)
                },
                "issues": issues,
                "suggestions": suggestions
            }, indent=2)
        }]
    }
            


@tool(
    "improve_resume_section",
    "Improve a specific section of the resume based on feedback. Returns guidance for improvement.",
    {
        "section": str,           # Section: "summary", "experience", "skills", "education", "header"
        "current_content": str,   # Current content (JSON string)
        "improvement_type": str,  # Type: "add_metrics", "shorten", "add_detail", "restructure"
        "feedback": str           # Specific feedback to address
    }
)
async def improve_resume_section(args: dict) -> dict:
    """
    Provide guidance for improving a specific resume section.

    Improvement types:
    - add_metrics: Add quantified achievements (numbers, percentages, dollar amounts)
    - shorten: Reduce length while preserving key information
    - add_detail: Expand with more specific information
    - restructure: Reorganize for better flow and impact
    """
    section = args["section"]
    current_content = args["current_content"]
    improvement_type = args["improvement_type"]
    feedback = args["feedback"]

    # Parse the current content
    try:
        content = json.loads(current_content)
    except json.JSONDecodeError:
        content = current_content

    suggestions = []
    examples = []

    if improvement_type == "add_metrics":
        suggestions = [
            "Add specific percentages (e.g., 'Increased efficiency by 35%')",
            "Include team/user numbers (e.g., 'Led team of 8 engineers')",
            "Quantify financial impact (e.g., 'Saved $50K annually')",
            "Add timeframes (e.g., 'Delivered 2 weeks ahead of schedule')",
            "Mention scale (e.g., 'Serving 10K+ daily active users')"
        ]
        examples = [
            "Before: 'Improved system performance' → After: 'Improved system performance by 40%, reducing load times from 3s to 1.8s'",
            "Before: 'Led development team' → After: 'Led team of 6 engineers delivering $2M product launch'",
            "Before: 'Managed customer accounts' → After: 'Managed portfolio of 50+ enterprise accounts worth $5M ARR'"
        ]

    elif improvement_type == "shorten":
        suggestions = [
            "Remove filler words (very, really, basically, actually)",
            "Start bullets with strong action verbs",
            "Combine related points into single impactful statements",
            "Keep each bullet under 100 characters",
            "Remove redundant adjectives"
        ]
        examples = [
            "Before: 'Successfully managed to implement a new system' → After: 'Implemented new system'",
            "Before: 'Was responsible for the development of' → After: 'Developed'",
            "Before: 'Worked closely with cross-functional teams to deliver' → After: 'Collaborated cross-functionally to deliver'"
        ]

    elif improvement_type == "add_detail":
        suggestions = [
            "Specify technologies, tools, and frameworks used",
            "Describe the business context and impact",
            "Mention scope and scale of projects",
            "Include relevant methodologies (Agile, Scrum, etc.)",
            "Add context about company size/industry"
        ]
        examples = [
            "Before: 'Built web application' → After: 'Built React/Node.js web application with PostgreSQL backend'",
            "Before: 'Improved processes' → After: 'Redesigned CI/CD pipeline using GitHub Actions, reducing deployment time'",
            "Before: 'Worked on backend' → After: 'Architected microservices using Python/FastAPI with Redis caching'"
        ]

    elif improvement_type == "restructure":
        suggestions = [
            "Lead with strongest, most relevant achievements",
            "Group related skills by category",
            "Use consistent verb tense (past for previous roles)",
            "Order experience by relevance, not just chronology",
            "Ensure parallel structure in bullet points"
        ]
        examples = [
            "Move leadership achievements to first bullet",
            "Group technical skills: Languages | Frameworks | Tools | Cloud",
            "Start all bullets with action verbs in same tense"
        ]

    return {
        "content": [{
            "type": "text",
            "text": json.dumps({
                "section": section,
                "improvement_type": improvement_type,
                "feedback_addressed": feedback,
                "suggestions": suggestions,
                "examples": examples,
                "instructions": f"Apply these improvements to the '{section}' section. After improving, use analyze_resume_data or update the resume JSON directly, then re-evaluate with evaluate_resume_quality."
            }, indent=2)
        }]
    }

# =============================================================================
# Hook Definition
# =============================================================================

async def log_tool_use(
    input_data: dict[str, Any],
    tool_use_id: str | None,
    context: HookContext
) -> dict[str, Any]:
    """Log all tool usage before execution."""
    tool_name = input_data.get('tool_name', 'unknown')
    print(f"[HOOK] Executing: {tool_name}")
    return {}


# =============================================================================
# MCP Server Setup
# =============================================================================

resume_tools = create_sdk_mcp_server(
    name="resume_tools",
    version="1.0.0",
    tools=[web_search, get_linkedin_profile, analyze_resume_data, generate_resume_docx, evaluate_resume_quality, improve_resume_section],
)


# =============================================================================
# Agent Runner
# =============================================================================

async def run_agent(args) -> None:
    """Run the resume generator agent."""
    output_path = args.output or f"{args.name.replace(' ', '_')}_Resume.docx"

    # Build allowed tools list - add quality tools if quality check is enabled
    allowed_tools = [
        "mcp__resume__web_search",
        "mcp__resume__get_linkedin_profile",
        "mcp__resume__analyze_resume_data",
        "mcp__resume__generate_resume_docx",
    ]

    if args.quality_check:
        allowed_tools.extend([
            "mcp__resume__evaluate_resume_quality",
            "mcp__resume__improve_resume_section",
        ])

    options = ClaudeAgentOptions(
        mcp_servers={"resume": resume_tools},
        allowed_tools=allowed_tools,
        hooks={
            'PreToolUse': [HookMatcher(hooks=[log_tool_use])]
        }
    )

    if args.quality_check:
        prompt = f"""Research {args.name} by searching the web and fetching their LinkedIn profile.
Then analyze the data and generate a professional resume document.

IMPORTANT: After generating the initial resume, you MUST run the evaluator-optimizer loop:

1. EVALUATE: Use evaluate_resume_quality with the resume JSON to get a score and feedback
2. CHECK: If score >= {args.min_score}, the resume passes. Save it and you're done.
3. IMPROVE: If score < {args.min_score}, use improve_resume_section to get improvement suggestions
4. APPLY: Apply the suggestions to create an improved resume JSON
5. REPEAT: Go back to step 1 and re-evaluate. Continue until score >= {args.min_score} or 3 iterations.

After the loop completes, save the final resume to: {output_path}

Always report the final score and any remaining suggestions."""
    else:
        prompt = (
            f"Research {args.name} by searching the web and fetching their LinkedIn profile. "
            f"Then analyze the data and generate a professional resume document. "
            f"Save the resume to: {output_path}"
        )

    async with ClaudeSDKClient(options=options) as client:
        await client.query(prompt)

        async for message in client.receive_response():
            if isinstance(message, AssistantMessage):
                for block in message.content:
                    if isinstance(block, TextBlock):
                        if args.verbose:
                            print(f"[ASSISTANT] {block.text}")
                    # if isinstance(block, ToolUseBlock):
                    #     if args.verbose:
                    #         print(f"[TOOL] {block.name}: {block.input}")

            if isinstance(message, ResultMessage):
                if args.verbose:
                    print(f"[RESULT] {message.result}")

    print(f"\nResume saved to: {output_path}")


# =============================================================================
# CLI Configuration
# =============================================================================

cli = AgentCLI(
    name="Resume Generator",
    description="Generate professional resumes using Claude Agent SDK with custom MCP tools",
    arguments=[
        CLIArgument(
            name="name",
            help="Name of the person to generate resume for"
        ),
        CLIArgument(
            name="--output",
            short="-o",
            help="Output file path (default: {name}_Resume.docx)"
        ),
        CLIArgument(
            name="--verbose",
            short="-v",
            help="Show detailed output including assistant messages",
            action="store_true"
        ),
        CLIArgument(
            name="--quality-check",
            short="-q",
            help="Evaluate the generated resume against quality criteria",
            action="store_true"
        ),
        CLIArgument(
            name="--min-score",
            short="-m",
            help="Minimum score required to pass quality check (default: 75)",
            arg_type=int,
            default=75,
        ),
    ],
    epilog="""
Examples:
  python resume_generator.py "John Smith"
  python resume_generator.py "Jane Doe" -o ./output/jane_resume.docx
  python resume_generator.py "John Smith" --verbose

Available mock profiles: John Smith, Jane Doe, Alex Johnson
"""
)


if __name__ == "__main__":
    exit_code = run_agent_cli(cli, run_agent)
    sys.exit(exit_code)
